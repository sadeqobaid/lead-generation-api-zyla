#!/usr/bin/env python3
"""Convert release Markdown documents to editable DOCX files.

Usage:
    python scripts/convert_markdown_to_docx.py --output-root dist/stage README.md scripts/README.md

Each output path mirrors the source path under ``--output-root`` and changes only the extension
from ``.md`` to ``.docx``. The converter intentionally handles the Markdown constructs used by
this project: headings, paragraphs, GFM tables, ordered/unordered lists, fenced code, blockquotes,
horizontal rules, links, inline code, and emphasis. It does not execute Markdown content.
"""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

INLINE_TOKEN = re.compile(r"(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|`.*?`|\[[^\]]+\]\([^\)]+\))")
TABLE_SEPARATOR = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def strip_markdown(text: str) -> str:
    """Return readable plain text without damaging identifiers such as ``JWT_SECRET``."""
    text = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1 (\2)", text)
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"(?<!\w)\*(.*?)\*(?!\w)", r"\1", text)
    text = re.sub(r"(?<!\w)_(.*?)_(?!\w)", r"\1", text)
    return html.unescape(text).strip()


def split_table_row(line: str) -> list[str]:
    """Split a pipe row while ignoring a single leading or trailing pipe."""
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [strip_markdown(cell) for cell in value.split("|")]


def add_inline_runs(paragraph, text: str) -> None:
    """Render the supported inline Markdown constructs as Word runs."""
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(html.unescape(text[cursor : match.start()]))
        token = match.group(0)
        run = paragraph.add_run()
        if token.startswith(("**", "__")):
            run.text = html.unescape(token[2:-2])
            run.bold = True
        elif token.startswith(("*", "_")):
            run.text = html.unescape(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run.text = token[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token)
            if link:
                run.text = f"{link.group(1)} ({link.group(2)})"
                run.underline = True
                run.font.color.rgb = None
            else:
                run.text = html.unescape(token)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(html.unescape(text[cursor:]))


def add_paragraph(document: Document, text: str, style: str | None = None, quote: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    if quote:
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        for run in paragraph.runs:
            run.italic = True
    add_inline_runs(paragraph, text)
    if quote:
        for run in paragraph.runs:
            run.italic = True


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.right_indent = Inches(0.3)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    for index, cell_text in enumerate(rows[0]):
        if index < width:
            table.rows[0].cells[index].text = cell_text
            for run in table.rows[0].cells[index].paragraphs[0].runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index in range(width):
            cells[index].text = row[index] if index < len(row) else ""


def configure_document(document: Document, source: Path) -> None:
    document.core_properties.author = "Sadeq Obaid"
    document.core_properties.last_modified_by = "Sadeq Obaid"
    document.core_properties.title = source.stem.replace("_", " ")
    document.core_properties.subject = "Lead Generation API Platform 0.1.0"
    document.core_properties.keywords = "Lead Generation API, Sadeq Obaid, commercial software"
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    for name, size in (("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
    if "No Spacing" not in [style.name for style in document.styles if style.type == WD_STYLE_TYPE.PARAGRAPH]:
        document.styles.add_style("No Spacing", WD_STYLE_TYPE.PARAGRAPH)


def convert(source: Path, destination: Path) -> None:
    document = Document()
    configure_document(document, source)
    lines = source.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph(document, " ".join(item.strip() for item in paragraph_buffer).strip())
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        if line.strip().startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(document, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_buffer.append(line)
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)), 3)
            add_paragraph(document, heading.group(2).strip(), style=f"Heading {level}")
            index += 1
            continue
        if line.strip() == "" :
            flush_paragraph()
            index += 1
            continue
        if line.strip().startswith(">"):
            flush_paragraph()
            add_paragraph(document, line.strip()[1:].strip(), quote=True)
            index += 1
            continue
        if re.match(r"^\s*([-*_])(?:\s*\1){2,}\s*$", line):
            flush_paragraph()
            document.add_paragraph("―" * 32)
            index += 1
            continue
        if re.match(r"^\s*[-+*]\s+", line):
            flush_paragraph()
            item = re.sub(r"^\s*[-+*]\s+", "", line)
            add_paragraph(document, item, style="List Bullet")
            index += 1
            continue
        if re.match(r"^\s*\d+[.)]\s+", line):
            flush_paragraph()
            item = re.sub(r"^\s*\d+[.)]\s+", "", line)
            add_paragraph(document, item, style="List Number")
            index += 1
            continue
        if "|" in line and index + 1 < len(lines) and TABLE_SEPARATOR.match(lines[index + 1]):
            flush_paragraph()
            rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, rows)
            continue
        paragraph_buffer.append(line)
        index += 1

    if in_code:
        add_code_block(document, code_buffer)
    flush_paragraph()
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-root", type=Path, required=True)
    parser.add_argument("sources", nargs="+", type=Path)
    args = parser.parse_args()
    for source in args.sources:
        destination = args.output_root / source.with_suffix(".docx")
        convert(source, destination)
        print(f"converted {source} -> {destination}")


if __name__ == "__main__":
    main()
